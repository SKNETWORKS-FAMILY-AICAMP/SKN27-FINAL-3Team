"""DOCX renderers for user-facing objection forms and analysis reports."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OFFICIAL_FINE_NOTICE_TITLE = "과태료 처분에 대한 이의신청서"


def render_report_docx(
    *,
    document_variant: str,
    title: str,
    form_data: dict[str, Any],
    sections: list[dict[str, Any]],
    petition_purpose: str,
    petition_reason: str,
) -> bytes:
    """Render one public document variant into an in-memory DOCX file."""

    document = Document()
    _set_default_font(document)
    if document_variant == "fine_notice":
        _render_fine_notice_form(
            document,
            form_data=form_data,
            petition_purpose=petition_purpose,
            petition_reason=petition_reason,
        )
    elif document_variant == "traffic_accident":
        _render_traffic_accident_form(document, title=title, form_data=form_data)
    else:
        _render_general_report(document, title=title, sections=sections)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _render_fine_notice_form(
    document: Document,
    *,
    form_data: dict[str, Any],
    petition_purpose: str,
    petition_reason: str,
) -> None:
    _add_title(document, OFFICIAL_FINE_NOTICE_TITLE)
    table = document.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    _fill_row(table, 0, "이의신청인 성명", _text(form_data.get("applicant_name")), "연락처", _text(form_data.get("contact")))
    _fill_row(table, 1, "주소", _text(form_data.get("address")), "차량번호", _text(form_data.get("vehicle_number")))
    _fill_row(table, 2, "수신 기관", _text(form_data.get("recipient")), "고지받은 일자", _text(form_data.get("notice_received_date")))
    _fill_row(table, 3, "과태료 금액", _text(form_data.get("fine_amount")), "고지서 번호", _text(form_data.get("case_number")))
    _fill_row(table, 4, "처분 사유", _text(form_data.get("violation_text")), "", "")
    _fill_row(
        table,
        5,
        "이의신청 내용",
        "\n\n".join(item for item in (petition_purpose, petition_reason) if item),
        "",
        "",
    )
    table.cell(5, 1).merge(table.cell(5, 3))

    document.add_paragraph()
    closing = document.add_paragraph("위 과태료 처분에 대하여 이의를 신청하오니 관계 법령에 따라 검토하여 주시기 바랍니다.")
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line = document.add_paragraph("년        월        일")
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signature = document.add_paragraph(f"신청인: {_text(form_data.get('applicant_name'))} (서명 또는 인)")
    signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"{_text(form_data.get('recipient'))} 귀하")


def _render_traffic_accident_form(document: Document, *, title: str, form_data: dict[str, Any]) -> None:
    _add_title(document, title or "교통사고 이의신청서")
    table = document.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    fields = (
        ("신청인", form_data.get("applicant_name")),
        ("수신 기관", form_data.get("recipient")),
        ("사건/접수 번호", form_data.get("case_number")),
        ("사고 일시·장소", _join_present(form_data.get("incident_at"), form_data.get("location"))),
        ("기존 조사 결과", form_data.get("investigation_result_summary")),
        ("이의신청 쟁점", form_data.get("objection_points")),
    )
    for row_index, (label, value) in enumerate(fields):
        _fill_cell(table.cell(row_index, 0), label, bold=True)
        _fill_cell(table.cell(row_index, 1), _text(value))

    _add_heading_and_body(document, "구체적 요청", _text(form_data.get("specific_request")))
    _add_heading_and_body(document, "제출 증빙", _text(form_data.get("evidence_summary")))


def _render_general_report(document: Document, *, title: str, sections: list[dict[str, Any]]) -> None:
    _add_title(document, title or "분석 리포트")
    for section in sections:
        if not isinstance(section, dict):
            continue
        heading = _text(section.get("title"))
        body = _text(section.get("body"))
        items = section.get("items")
        if heading:
            _add_heading_and_body(document, heading, body)
        elif body:
            document.add_paragraph(body)
        if isinstance(items, list):
            for item in items:
                text = _text(item)
                if text:
                    document.add_paragraph(text, style="List Bullet")


def _add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def _add_heading_and_body(document: Document, heading: str, body: str) -> None:
    if heading:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(heading)
        run.bold = True
        run.font.size = Pt(12)
    if body:
        document.add_paragraph(body)


def _fill_row(table, row_index: int, first_label: str, first_value: str, second_label: str, second_value: str) -> None:
    _fill_cell(table.cell(row_index, 0), first_label, bold=True)
    _fill_cell(table.cell(row_index, 1), first_value)
    _fill_cell(table.cell(row_index, 2), second_label, bold=True)
    _fill_cell(table.cell(row_index, 3), second_value)


def _fill_cell(cell, value: str, *, bold: bool = False) -> None:
    cell.text = value
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.bold = bold


def _set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), "맑은 고딕")


def _set_table_borders(table) -> None:
    table.autofit = False
    for column in table.columns:
        column.width = Cm(3.6)
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = borders.makeelement(qn(f"w:{edge}"), {})
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), "000000")


def _join_present(*values: Any) -> str:
    return " / ".join(text for value in values if (text := _text(value)))


def _text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return " / ".join(_text(item) for item in value if _text(item))
    return str(value).strip()
