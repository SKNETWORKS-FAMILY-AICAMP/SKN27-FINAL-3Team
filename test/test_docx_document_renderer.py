from __future__ import annotations

from io import BytesIO

from docx import Document

from ai.agents import objection_report_generation as objection_reporting


def _render(document_variant: str, **payload):
    renderer = getattr(objection_reporting, "render_report_docx", None)
    assert callable(renderer)
    return renderer(document_variant=document_variant, **payload)


def _document_text(body: bytes) -> str:
    document = Document(BytesIO(body))
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def test_fine_notice_renderer_creates_official_docx_without_resident_id():
    body = _render(
        "fine_notice",
        title="과태료 부과 처분 이의신청서",
        form_data={
            "applicant_name": "홍길동",
            "resident_registration_number": "900101-1234567",
            "address": "서울특별시 강남구 테헤란로 123",
            "contact": "010-1234-5678",
            "vehicle_number": "12가3456",
            "recipient": "강남구청",
            "notice_received_date": "2026-06-10",
            "fine_amount": "70,000원",
            "case_number": "2026-001",
            "violation_text": "신호 위반",
        },
        sections=[],
        petition_purpose="처분의 취소 또는 감경을 요청합니다.",
        petition_reason="고지서 기재 사실관계와 실제 상황이 다릅니다.",
    )

    assert body[:2] == b"PK"
    text = _document_text(body)
    assert "과태료 처분에 대한 이의신청서" in text
    assert "처분의 취소 또는 감경을 요청합니다." in text
    assert "900101-1234567" not in text


def test_traffic_accident_renderer_uses_existing_form_data():
    body = _render(
        "traffic_accident",
        title="교통사고 이의신청서 초안",
        form_data={
            "applicant_name": "김운전자",
            "recipient": "강남경찰서",
            "case_number": "ACC-2026-1",
            "incident_at": "2026-06-10 13:33",
            "location": "서울시 강남구 교차로",
            "investigation_result_summary": "기존 과실비율 70:30",
            "objection_points": ["선진입", "영상 누락"],
            "specific_request": "블랙박스 영상을 재검토해 주세요.",
        },
        sections=[],
        petition_purpose="",
        petition_reason="",
    )

    assert body[:2] == b"PK"
    text = _document_text(body)
    assert "김운전자" in text
    assert "강남경찰서" in text
    assert "선진입" in text
    assert "블랙박스 영상을 재검토해 주세요." in text


def test_general_report_renderer_uses_report_title_and_sections():
    body = _render(
        "general",
        title="분석 리포트",
        form_data={},
        sections=[
            {"title": "사실관계", "body": "신청인이 제출한 사실 요약"},
            {"title": "한계", "body": "추가 증빙 확인 필요"},
        ],
        petition_purpose="",
        petition_reason="",
    )

    assert body[:2] == b"PK"
    text = _document_text(body)
    assert "분석 리포트" in text
    assert "신청인이 제출한 사실 요약" in text
    assert "추가 증빙 확인 필요" in text
