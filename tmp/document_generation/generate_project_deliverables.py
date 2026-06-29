from __future__ import annotations

import html
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph as PdfParagraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
PROJECT_TEMPLATE = Path(r"C:\Users\Playdata\Downloads\[기획] 프로젝트 기획서_양식_27기_0팀.docx")
DATA_TEMPLATE = Path(r"C:\Users\Playdata\Downloads\[데이터 수집 및 저장] 수집 데이터 보고서_27기_0팀.docx")
DOCX_OUT = ROOT / "output" / "docx"
PDF_OUT = ROOT / "output" / "pdf"
TMP_OUT = ROOT / "tmp" / "document_generation"

PROJECT_DOCX = DOCX_OUT / "프로젝트_기획서_27기_3팀_2026-06-22.docx"
DATA_DOCX = DOCX_OUT / "수집_데이터_보고서_27기_3팀_2026-06-22.docx"
PROJECT_PDF = PDF_OUT / "프로젝트_기획서_27기_3팀_2026-06-22.pdf"
DATA_PDF = PDF_OUT / "수집_데이터_보고서_27기_3팀_2026-06-22.pdf"
SCREEN_PDF = PDF_OUT / "화면설계서_27기_3팀_2026-06-22.pdf"

PROJECT_NAME = "교통분쟁 AI: 과실비율·범칙금/과태료 분석 및 리포팅 서비스"
TEAM_LABEL = "SK 네트웍스 Family AI 27기 : 3팀"
GITHUB_URL = "https://github.com/SKNETWORKS-FAMILY-AICAMP/SKN27-FINAL-3Team"
WRITE_DATE = "2026. 6. 22."
AUTHORS = "3팀: hi20260204-maker, leejaegang27, ohjuheecode, techshin31, workzion2"


FONT_CANDIDATES = [
    (
        "NotoSansKR",
        Path("C:/Windows/Fonts/NotoSansKR-Regular.ttf"),
        Path("C:/Windows/Fonts/NotoSansKR-Bold.ttf"),
    ),
    (
        "MalgunGothic",
        Path("C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/malgunbd.ttf"),
    ),
]


def register_pdf_font() -> str:
    for family, regular, bold in FONT_CANDIDATES:
        if regular.exists() and bold.exists():
            pdfmetrics.registerFont(TTFont(family, str(regular)))
            pdfmetrics.registerFont(TTFont(f"{family}-Bold", str(bold)))
            pdfmetrics.registerFontFamily(
                family,
                normal=family,
                bold=f"{family}-Bold",
                italic=family,
                boldItalic=f"{family}-Bold",
            )
            return family
    raise RuntimeError("한글 PDF 생성을 위한 TTF 폰트를 찾지 못했습니다.")


PDF_FONT = register_pdf_font()


def ensure_dirs() -> None:
    DOCX_OUT.mkdir(parents=True, exist_ok=True)
    PDF_OUT.mkdir(parents=True, exist_ok=True)
    TMP_OUT.mkdir(parents=True, exist_ok=True)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def apply_doc_defaults(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    for style_name, size, bold, color in [
        ("Normal", 10, False, "111827"),
        ("Heading 1", 15, True, "1F4E79"),
        ("Heading 2", 12, True, "1F4E79"),
        ("List Bullet", 9.5, False, "111827"),
    ]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "맑은 고딕"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            style.font.size = Pt(size)
            style.font.bold = bold
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_after = Pt(5)
            style.paragraph_format.line_spacing = 1.12


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def set_paragraph_text(paragraph, text: str, size: float = 10, bold: bool = False, color: str = "111827") -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def set_cell_text(cell, text: str, size: float = 9.3, bold: bool = False, color: str = "111827") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def style_table(table, header_rows: int = 1) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, size=8.5 if row_index >= header_rows else 8.7, bold=row_index < header_rows)
            if row_index < header_rows:
                shade_cell(cell, "1F4E79")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size=8.7, bold=True, color="FFFFFF")
            elif row_index % 2 == 0:
                shade_cell(cell, "F8FAFC")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        run = new_paragraph.add_run(text)
        set_run_font(run, size=9.7)
    new_paragraph.paragraph_format.space_after = Pt(4)
    new_paragraph.paragraph_format.line_spacing = 1.12
    return new_paragraph


def insert_bullets_after(anchor: Paragraph, items: list[str]) -> Paragraph:
    current = anchor
    for item in items:
        current = insert_paragraph_after(current, item, "List Bullet")
        for run in current.runs:
            set_run_font(run, size=9.4)
    return current


def insert_paragraphs_after(anchor: Paragraph, paragraphs: list[str]) -> Paragraph:
    current = anchor
    for item in paragraphs:
        current = insert_paragraph_after(current, item, "Normal")
    return current


def fill_common_template_tables(doc: Document, report_title: str, stage: str) -> None:
    if doc.tables:
        for cell in doc.tables[0].rows[0].cells:
            set_cell_text(cell, f"{TEAM_LABEL} {report_title}", size=11, bold=True, color="FFFFFF")
            shade_cell(cell, "1F4E79")
    if len(doc.tables) > 1:
        table = doc.tables[1]
        values = {
            "산출물 단계": stage,
            "제출 일자": WRITE_DATE,
            "깃허브 경로": GITHUB_URL,
            "작성 팀원": AUTHORS,
        }
        for row in table.rows:
            key = row.cells[0].text.strip()
            if key in values:
                set_cell_text(row.cells[1], values[key], size=9.3)
        style_table(table, header_rows=0)


def build_project_plan() -> Path:
    doc = Document(PROJECT_TEMPLATE)
    apply_doc_defaults(doc)
    fill_common_template_tables(doc, "프로젝트 기획서", "기획")

    if len(doc.tables) > 2:
        for cell in doc.tables[2].rows[0].cells:
            set_cell_text(cell, "프로젝트 기획", size=12, bold=True, color="FFFFFF")
            shade_cell(cell, "1F4E79")

    sections = {
        "프로젝트 주제": [
            f"프로젝트명은 '{PROJECT_NAME}'다.",
            "MVP는 사용자가 고지서, 사고 설명, 사진 또는 영상을 입력했을 때 과태료/범칙금 분석, 과실비율 관련 쟁점 정리, 근거 기반 답변, 리포트 및 이의신청서 초안으로 이어지는 사용자 흐름을 제공하는 것을 목표로 한다.",
            "현재 저장소는 기능 구현 완료 단계가 아니라 요구사항, WBS, Agent 결과 Schema, 화면설계, 통합 시나리오를 정리하는 초기 설계 단계로 확인된다.",
        ],
        "문제 정의": [
            "교통사고와 과태료/범칙금 상황에서 사용자는 고지서 필드, 납부기한, 처분 단계, 이의제기 가능성, 사고 쟁점, 필요한 증거를 한 번에 이해하기 어렵다.",
            "단순 챗봇 답변은 법령 근거, 유사 사례, OCR 신뢰도, 자료 부족 여부를 구조적으로 남기기 어렵다.",
            "따라서 본 프로젝트는 AI 답변을 최종 판단처럼 단정하지 않고, 근거와 한계, 다음 행동을 함께 제공하는 리포팅 서비스로 설계한다.",
        ],
        "시장조사 및 BM 분석": [
            "회의 문서 기준으로 국내 교통 민원, 해외 과태료 이의제기, OCR 리포팅, 이미지/영상 분석, 텍스트 입력 분석 서비스를 벤치마킹 대상으로 조사하기로 했다.",
            "현재 저장소에는 특정 서비스별 확정 비교표가 없으므로, 본 기획서에서는 벤치마킹 결론을 완료 항목으로 쓰지 않는다.",
            "도입할 BM 포인트는 구조화 입력, 근거 기반 답변, 결과 리포트 저장, PDF/DOCX 다운로드, 법률 단정 방지, 자료 부족 시 추가 질문 흐름이다.",
        ],
        "시스템 구성 기획": [
            "사용자 흐름은 로그인/서비스 설명 진입 -> AI 교통 상담 챗봇 -> Supervisor 입력 분류 -> 고지서 OCR/법률 근거/텍스트 ML/영상 이미지/이의신청서 생성 노드 -> 결과 카드와 리포트 -> 마이페이지 저장 순서로 잡는다.",
            "문서화된 폴더 책임은 app(화면/API), ai(Supervisor/Agent/Schema), etl(데이터 수집/전처리), storage(DB/RAG/파일 저장), test(단위/통합/E2E/수동 시나리오), docs(source of truth)로 분리된다.",
            "Agent 결과는 node_name, node_code, status, summary, structured_result, evidence, next_actions, limitations를 공통 envelope으로 반환하고, Supervisor가 최종 답변을 병합한다.",
        ],
        "모델링 계획": [
            "고지서 OCR·과태료/범칙금 분석 노드는 OCR 결과, 고지 정보, 처분 단계, 이의제기 가능성, 부족 서류, 필요 증거를 구조화한다.",
            "법률 근거 검색 노드는 도로교통법, 시행령, 시행규칙, 행정 기준, 고시 등 법률 계열 데이터를 검색하고 조문 metadata와 적용 한계를 반환한다.",
            "텍스트 ML/판례·사례 검색 노드는 경위서, OCR 텍스트, 판례, 유튜브 자막, 과실비율심의사례를 청크/임베딩/요약/태그 형태로 다룬다.",
            "영상·이미지 분석 노드는 key frame, 장면 요약, detected object, confidence, 품질 이슈를 반환하되 사고 책임이나 과실비율을 확정하지 않는다.",
            "회의 기준으로 모델은 확정된 상태가 아니며, 최소 2개 이상 모델 비교와 샘플 검증 이후 채택 여부를 결정한다.",
        ],
        "데이터 수집 전략": [
            "법률 데이터는 도로교통법, 시행령, 시행규칙, 과태료/범칙금 관련 행정 기준과 고시를 중심으로 수집한다.",
            "과태료·범칙금 분석용 룰/매핑 데이터는 법률 원문 DB와 분리하고, 위반 유형, 금액, 벌점, 예외 조건, 처분 단계 판단에 필요한 구조화 데이터로 관리한다.",
            "과실비율 영역은 판례, 유튜브 자막 사고 사례, 과실비율심의사례를 RAG/ML 입력으로 정제한다.",
            "영상·이미지 영역은 비식별 샘플, frame metadata, key frame, scene summary 중심으로 manifest를 구성한다.",
            "원본 고지서, 블랙박스, 개인정보 포함 파일은 기본 커밋 대상에서 제외하고 UTF-8, metadata, source_type, domain 기준으로 추적한다.",
        ],
        "역할분담(R&R)": [
            "hi20260204-maker: WBS/문서, Supervisor 통합 답변 구조, 홈·로그인·챗봇 진입 흐름, 이의신청서 생성 노드, 통합 QA.",
            "leejaegang27: 경위서/OCR 결과 처리, 텍스트 ML, 과실비율 판례, 유튜브 자막 사례, 과실비율심의사례 데이터, 판례 Agent.",
            "ohjuheecode: 차량 사고 이미지·영상 데이터셋, Vision/DL 분석, 영상·이미지 Agent, DL 결과 구조화.",
            "techshin31: 법률 데이터 수집, 전처리, DB 적재, 법률 원문/조문/근거 metadata.",
            "workzion2: 고지서 OCR, 과태료·범칙금·벌칙 분석용 룰/매핑 데이터, 과태료·범칙금 분석 흐름.",
        ],
    }

    for paragraph in list(doc.paragraphs):
        heading = paragraph.text.strip()
        if heading in sections:
            for run in paragraph.runs:
                set_run_font(run, size=12, bold=True, color="1F4E79")
            insert_paragraphs_after(paragraph, sections[heading])

    doc.save(PROJECT_DOCX)
    return PROJECT_DOCX


def ensure_table_rows(table, row_count: int) -> None:
    while len(table.rows) < row_count:
        table.add_row()


def fill_rows(table, rows: list[list[str]], start: int = 0) -> None:
    ensure_table_rows(table, start + len(rows))
    for row_offset, values in enumerate(rows):
        row = table.rows[start + row_offset]
        for index, value in enumerate(values):
            if index < len(row.cells):
                set_cell_text(row.cells[index], value, size=8.2, bold=start + row_offset == 0)
    style_table(table, header_rows=1 if start == 0 else 0)


def build_data_report() -> Path:
    doc = Document(DATA_TEMPLATE)
    apply_doc_defaults(doc)
    fill_common_template_tables(doc, "수집 데이터 보고서", "데이터 수집 및 저장")

    if len(doc.tables) > 2:
        for cell in doc.tables[2].rows[0].cells:
            set_cell_text(cell, "수집 데이터", size=12, bold=True, color="FFFFFF")
            shade_cell(cell, "1F4E79")

    dataset_rows = [
        ["데이터명", "수집 대상", "수집 목적", "사용 예정 기능", "출처 / 저작권"],
        ["법률 원문 데이터", "도로교통법, 시행령, 시행규칙, 행정 기준, 고시", "과태료·범칙금 분석과 법률 근거 검색", "법률 근거 검색 노드, RAG evidence", "공개 법령/API, 출처 metadata 필요"],
        ["과태료·범칙금 룰/매핑", "위반 유형, 금액, 벌점, 예외 조건, 처분 단계", "고지서 분석 결과 구조화", "고지서 OCR·과태료/범칙금 분석 노드", "법령 기반 내부 구조화 데이터"],
        ["판례·심의·사고 사례 텍스트", "판례, 과실비율심의사례, 유튜브 자막 사고 사례", "사고 유형, 쟁점, 유사 사례 검색", "텍스트 ML/판례·사례 검색 노드", "공개 원문/자막, 사용권 검토 필요"],
        ["영상·이미지 manifest", "사고 사진, 차량 파손 이미지, 블랙박스/사고 영상, key frame", "Vision/DL 분석과 장면 요약", "영상·이미지 분석 노드", "비식별 샘플만 저장, 원본 민감 파일 커밋 금지"],
        ["사용자 업로드 자료", "고지서, 경위서, 사고 사진, 영상, 사용자 설명", "개별 사건 분석 입력", "챗봇, 리포트, 이의신청서 초안", "사용자 제공 자료, 개인정보 보호 대상"],
    ]
    fill_rows(doc.tables[3], dataset_rows)

    method_rows = [
        ["수집 방식", "선택", "수집 방식", "선택"],
        ["웹 크롤링", "부분 사용", "API 호출", "사용"],
        ["사용자 입력", "사용", "문서 파일 업로드(PDF, 이미지, 영상 등)", "사용"],
        ["기타: 회의/이슈/요구사항 문서 기반 데이터 범위 정리", "사용", "", ""],
    ]
    fill_rows(doc.tables[4], method_rows)

    tech_rows = [
        ["사용 언어 / 라이브러리", "Python requests/BeautifulSoup, OCR 도구 후보, RAG chunking/embedding 도구, Django API. 실제 라이브러리는 구현 문서 확정 후 고정."],
        ["자동화 여부 및 주기", "현재는 설계 단계다. 법률/룰/사례/영상 manifest는 ETL 파이프라인으로 자동화 예정이며, 샘플 검증 전까지는 수동 검토를 병행한다."],
        ["오류 발생 시 예외 처리 전략", "OCR 실패, 필수 필드 누락, RAG 검색 실패, Agent 응답 실패, 파일 업로드 실패를 status=partial/failed와 limitations로 반환한다."],
    ]
    fill_rows(doc.tables[5], tech_rows)

    field_rows = [
        ["파일명 / 테이블명", "필드명", "데이터 타입", "설명", "예시"],
        ["legal_documents", "law_name", "string", "법령명", "도로교통법"],
        ["legal_documents", "article", "string", "조문 번호", "제32조"],
        ["legal_documents", "effective_date", "date", "시행일 또는 기준일", "2026-06-22"],
        ["fine_rules", "violation_type", "string", "위반 유형", "어린이보호구역 주정차 위반"],
        ["fine_rules", "fine_amount", "integer", "과태료 또는 범칙금 금액", "120000"],
        ["fault_cases", "issue_tags", "array", "과실 쟁점 태그", "신호, 좌회전, 선진입"],
        ["media_manifest", "file_type", "string", "이미지/영상 유형", "blackbox_video"],
        ["agent_results", "evidence", "array", "근거 목록", "law, precedent, vision_result"],
    ]
    fill_rows(doc.tables[6], field_rows)

    amount_rows = [
        ["전체 수집 데이터 건수", "문서 기준 법령/시행령/시행규칙 총 23개 수집 완료로 기록됨. 저장소 내 raw 데이터 파일은 확인되지 않아 원본 적재 건수는 검증 필요."],
        ["추출된 고품질 데이터 건수 (필터링 후 기준)", "법률 데이터 외 판례, 유튜브 자막, 과실비율심의사례, 영상·이미지 샘플은 진행 중 또는 검증 필요 상태로 기록한다."],
    ]
    fill_rows(doc.tables[7], amount_rows)

    storage_rows = [
        ["저장 경로", "설계 기준: etl/legal, etl/fine_rules, etl/fault_cases, etl/vision_manifest, storage/rag, storage/samples. 현재 raw 데이터 저장 경로는 미확정."],
        ["저장 포맷", "CSV, JSON, DB table, RAG document/chunk 모두 후보. 구현 전 schema 확정 필요."],
        ["인코딩", "UTF-8 고정. HTML, Markdown, CSV/JSON, DB 입출력 모두 한글 깨짐 방지 기준 적용."],
    ]
    fill_rows(doc.tables[8], storage_rows)

    legal_rows = [
        ["개인정보 포함 여부", "법령/공개 사례 데이터는 미포함 가능성이 높으나, 사용자 업로드 고지서/영상/사진은 개인정보 또는 민감정보 포함 가능."],
        ["포함된 경우 필드", "이름, 차량번호, 주소, 연락처, 위반 일시/장소, 영상/사진 속 개인 식별 정보."],
        ["비식별화 조치 여부", "원본 민감 데이터는 Git 커밋 금지. 테스트 fixture는 비식별화 또는 샘플 대체 후 사용."],
        ["출처 및 사용권 / 공개 여부", "공개 법령은 출처와 최신성 metadata를 남긴다. 판례/자막/영상은 원문 링크와 사용권 검토 필요."],
        ["라이선스 / 약관 검토 여부", "진행 중. 유튜브 자막, 보험사 자료, 영상 데이터는 사용 범위 확인 필요."],
        ["검토자", "3팀 공동 검토"],
        ["검토 일자", WRITE_DATE],
    ]
    fill_rows(doc.tables[9], legal_rows)

    quality_rows = [
        ["중복 제거 기준", "source_url, law_name+article, case_id, video_id+timestamp_range, file_hash 기준 중복 제거."],
        ["정합성 검증 방법", "필수 필드 누락, 날짜 형식, 금액/벌점 타입, source_type/domain, evidence metadata를 검증."],
        ["Null 처리 및 결측치 전략", "필수 입력이 없으면 partial 상태와 missing_fields를 반환하고, 분석보다 추가 질문을 우선한다."],
        ["표준화 전략", "HTML 제거, 광고/불필요 문구 제거, 문단 분리, chunk 분할, metadata 정규화, UTF-8 저장."],
    ]
    fill_rows(doc.tables[10], quality_rows)

    history_rows = [
        ["변경일", "변경자", "변경 내용", "비고"],
        ["2026-06-17", "3팀", "요구사항 정의서 v0.4 업데이트", "125개 요구사항으로 재정리"],
        ["2026-06-18", "3팀", "WBS, 역할, Agent/Supervisor 구조 정리", "중간/최종 일정 기준 반영"],
        ["2026-06-22", "3팀", "수집 데이터 보고서 PDF 작성", "raw 데이터 미확인 항목은 검증 필요로 표기"],
    ]
    fill_rows(doc.tables[11], history_rows)

    flow_anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "(이미지, 순서도 또는 코드)":
            flow_anchor = paragraph
            break
    if flow_anchor is not None:
        set_paragraph_text(flow_anchor, "수집/저장 흐름")
        insert_paragraphs_after(
            flow_anchor,
            [
                "1. 원천 데이터 확인: 법령/API, 과태료·범칙금 기준, 판례/심의사례, 자막, 영상·이미지 샘플, 사용자 업로드 자료를 구분한다.",
                "2. 수집 및 전처리: HTML 제거, 중복 제거, OCR/텍스트 정제, frame/key scene 추출, metadata 부여를 수행한다.",
                "3. 저장 및 색인: DB table, RAG document/chunk, 파일 manifest로 분리 저장한다.",
                "4. 검증 및 활용: Agent 결과 envelope의 evidence, limitations, next_actions에 연결한다.",
            ],
        )

    doc.save(DATA_DOCX)
    return DATA_DOCX


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    common = {"fontName": PDF_FONT, "wordWrap": "CJK", "splitLongWords": True}
    return {
        "title": ParagraphStyle("TitleKo", parent=base["Title"], fontName=f"{PDF_FONT}-Bold", fontSize=22, leading=30, alignment=TA_CENTER, textColor=colors.HexColor("#111827"), spaceAfter=6 * mm),
        "subtitle": ParagraphStyle("SubtitleKo", parent=base["Normal"], **common, fontSize=9.5, leading=14, alignment=TA_CENTER, textColor=colors.HexColor("#4B5563"), spaceAfter=8 * mm),
        "h1": ParagraphStyle("H1Ko", parent=base["Heading1"], fontName=f"{PDF_FONT}-Bold", fontSize=14, leading=20, textColor=colors.HexColor("#111827"), spaceBefore=5 * mm, spaceAfter=2.5 * mm, keepWithNext=True, wordWrap="CJK"),
        "h2": ParagraphStyle("H2Ko", parent=base["Heading2"], fontName=f"{PDF_FONT}-Bold", fontSize=11.2, leading=16, textColor=colors.HexColor("#1F2937"), spaceBefore=3 * mm, spaceAfter=2 * mm, keepWithNext=True, wordWrap="CJK"),
        "body": ParagraphStyle("BodyKo", parent=base["BodyText"], **common, fontSize=8.7, leading=13, textColor=colors.HexColor("#111827"), spaceAfter=1.8 * mm),
        "small": ParagraphStyle("SmallKo", parent=base["BodyText"], **common, fontSize=7.2, leading=10.5, textColor=colors.HexColor("#4B5563"), spaceAfter=1 * mm),
        "thead": ParagraphStyle("HeadKo", parent=base["BodyText"], fontName=f"{PDF_FONT}-Bold", fontSize=6.8, leading=9.5, textColor=colors.white, alignment=TA_CENTER, wordWrap="CJK"),
        "tbody": ParagraphStyle("CellKo", parent=base["BodyText"], **common, fontSize=6.6, leading=9.2, textColor=colors.HexColor("#111827")),
    }


PDF_STYLES = pdf_styles()


def esc(text: object) -> str:
    return html.escape(str(text)).replace("\n", "<br/>")


def pp(text: object, style: str = "body") -> PdfParagraph:
    return PdfParagraph(esc(text), PDF_STYLES[style])


def ptable(rows: list[list[object]], widths: list[float], header: bool = True) -> Table:
    converted = []
    for index, row in enumerate(rows):
        style = "thead" if header and index == 0 else "tbody"
        converted.append([pp(cell, style) for cell in row])
    table = Table(converted, colWidths=widths, repeatRows=1 if header else 0, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79") if header else colors.white),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D1D5DB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
            ]
        )
    )
    return table


def optimize_image(path: Path) -> Path:
    cache_dir = TMP_OUT / "optimized-images"
    cache_dir.mkdir(parents=True, exist_ok=True)
    out = cache_dir / f"{path.stem}-screen-pdf.jpg"
    if out.exists() and out.stat().st_mtime >= path.stat().st_mtime:
        return out
    with PILImage.open(path) as image:
        if image.mode in ("RGBA", "LA"):
            background = PILImage.new("RGB", image.size, "white")
            background.paste(image.convert("RGBA"), mask=image.getchannel("A"))
            image = background
        else:
            image = image.convert("RGB")
        max_width = 1800
        if image.width > max_width:
            ratio = max_width / image.width
            image = image.resize((max_width, max(1, int(image.height * ratio))), PILImage.Resampling.LANCZOS)
        image.save(out, "JPEG", quality=78, optimize=True, progressive=True)
    return out


def add_pdf_image(story: list, image_path: Path, caption: str, max_width: float, max_height: float) -> None:
    if not image_path.exists():
        story.append(pp(f"이미지 파일 없음: {image_path}", "small"))
        return
    optimized = optimize_image(image_path)
    with PILImage.open(optimized) as image:
        width, height = image.size
    scale = min(max_width / width, max_height / height)
    story.append(Image(str(optimized), width=width * scale, height=height * scale))
    story.append(pp(caption, "small"))
    story.append(Spacer(1, 2 * mm))


def portrait_footer(label: str):
    def _footer(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont(PDF_FONT, 7)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawString(doc.leftMargin, 10 * mm, label)
        canvas.drawRightString(A4[0] - doc.rightMargin, 10 * mm, str(doc.page))
        canvas.restoreState()

    return _footer


def landscape_footer(label: str):
    def _footer(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont(PDF_FONT, 7)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawString(doc.leftMargin, 8 * mm, label)
        canvas.drawRightString(landscape(A4)[0] - doc.rightMargin, 8 * mm, str(doc.page))
        canvas.restoreState()

    return _footer


def add_section(story: list, title: str, paragraphs: list[str]) -> None:
    story.append(pp(title, "h1"))
    for paragraph in paragraphs:
        story.append(pp(paragraph))


def build_project_plan_pdf() -> Path:
    story: list = [
        pp("프로젝트 기획서", "title"),
        pp(f"{TEAM_LABEL}\n프로젝트명: {PROJECT_NAME}\n작성일: {WRITE_DATE}", "subtitle"),
        ptable(
            [
                ["항목", "내용"],
                ["산출물 단계", "기획"],
                ["깃허브 경로", GITHUB_URL],
                ["작성 팀원", AUTHORS],
                ["작성 기준", "요구사항 정의서 v0.4, WBS/역할 문서, Agent 결과 Schema, 화면설계서, 2026-06-18/06-22 회의 문서"],
            ],
            [35 * mm, 136 * mm],
        ),
    ]

    add_section(
        story,
        "1. 프로젝트 주제",
        [
            f"프로젝트명은 '{PROJECT_NAME}'다.",
            "MVP는 사용자가 고지서, 사고 설명, 사진 또는 영상을 입력했을 때 과태료/범칙금 분석, 과실비율 관련 쟁점 정리, 근거 기반 답변, 리포트 및 이의신청서 초안으로 이어지는 사용자 흐름을 제공하는 것을 목표로 한다.",
            "현재 저장소는 기능 구현 완료 단계가 아니라 요구사항, WBS, Agent 결과 Schema, 화면설계, 통합 시나리오를 정리하는 초기 설계 단계로 확인된다.",
        ],
    )
    add_section(
        story,
        "2. 문제 정의",
        [
            "교통사고와 과태료/범칙금 상황에서 사용자는 고지서 필드, 납부기한, 처분 단계, 이의제기 가능성, 사고 쟁점, 필요한 증거를 한 번에 이해하기 어렵다.",
            "단순 챗봇 답변은 법령 근거, 유사 사례, OCR 신뢰도, 자료 부족 여부를 구조적으로 남기기 어렵다.",
            "따라서 본 프로젝트는 AI 답변을 최종 판단처럼 단정하지 않고, 근거와 한계, 다음 행동을 함께 제공하는 리포팅 서비스로 설계한다.",
        ],
    )
    add_section(
        story,
        "3. 시장조사 및 BM 분석",
        [
            "회의 문서 기준으로 국내 교통 민원, 해외 과태료 이의제기, OCR 리포팅, 이미지/영상 분석, 텍스트 입력 분석 서비스를 벤치마킹 대상으로 조사하기로 했다.",
            "현재 저장소에는 특정 서비스별 확정 비교표가 없으므로, 본 기획서에서는 벤치마킹 결론을 완료 항목으로 쓰지 않는다.",
            "도입할 BM 포인트는 구조화 입력, 근거 기반 답변, 결과 리포트 저장, PDF/DOCX 다운로드, 법률 단정 방지, 자료 부족 시 추가 질문 흐름이다.",
        ],
    )
    add_section(
        story,
        "4. 시스템 구성 기획",
        [
            "사용자 흐름은 로그인/서비스 설명 진입 -> AI 교통 상담 챗봇 -> Supervisor 입력 분류 -> 고지서 OCR/법률 근거/텍스트 ML/영상 이미지/이의신청서 생성 노드 -> 결과 카드와 리포트 -> 마이페이지 저장 순서로 잡는다.",
            "문서화된 폴더 책임은 app(화면/API), ai(Supervisor/Agent/Schema), etl(데이터 수집/전처리), storage(DB/RAG/파일 저장), test(단위/통합/E2E/수동 시나리오), docs(source of truth)로 분리된다.",
            "Agent 결과는 node_name, node_code, status, summary, structured_result, evidence, next_actions, limitations를 공통 envelope으로 반환하고, Supervisor가 최종 답변을 병합한다.",
        ],
    )
    story.append(pp("5. 모델링 계획", "h1"))
    story.append(
        ptable(
            [
                ["노드", "역할", "현재 상태"],
                ["고지서 OCR·과태료/범칙금 분석 노드", "OCR 결과, 고지 정보, 처분 단계, 이의제기 가능성, 부족 서류, 필요 증거 구조화", "필드와 샘플 검증 필요"],
                ["법률 근거 검색 노드", "도로교통법, 시행령, 시행규칙, 행정 기준, 고시 검색 및 metadata 반환", "법률 source/schema 보완 필요"],
                ["텍스트 ML/판례·사례 검색 노드", "경위서, OCR 텍스트, 판례, 유튜브 자막, 심의사례 청크/임베딩/요약/태그", "모델 후보와 데이터 수집 진행 중"],
                ["영상·이미지 분석 노드", "key frame, 장면 요약, detected object, confidence, 품질 이슈 반환", "POC 샘플 검증 필요"],
                ["이의신청서 생성/리포트 노드", "분석 결과와 법률 근거를 받아 초안/리포트 구조 생성", "사실관계 부족 시 추가 질문 필요"],
            ],
            [48 * mm, 83 * mm, 40 * mm],
        )
    )
    add_section(
        story,
        "6. 데이터 수집 전략",
        [
            "법률 데이터는 도로교통법, 시행령, 시행규칙, 과태료/범칙금 관련 행정 기준과 고시를 중심으로 수집한다.",
            "과태료·범칙금 분석용 룰/매핑 데이터는 법률 원문 DB와 분리하고, 위반 유형, 금액, 벌점, 예외 조건, 처분 단계 판단에 필요한 구조화 데이터로 관리한다.",
            "과실비율 영역은 판례, 유튜브 자막 사고 사례, 과실비율심의사례를 RAG/ML 입력으로 정제한다.",
            "영상·이미지 영역은 비식별 샘플, frame metadata, key frame, scene summary 중심으로 manifest를 구성한다.",
            "원본 고지서, 블랙박스, 개인정보 포함 파일은 기본 커밋 대상에서 제외하고 UTF-8, metadata, source_type, domain 기준으로 추적한다.",
        ],
    )
    story.append(pp("7. 역할분담(R&R)", "h1"))
    story.append(
        ptable(
            [
                ["담당", "주요 책임"],
                ["hi20260204-maker", "WBS/문서, Supervisor 통합 답변 구조, 홈·로그인·챗봇 진입 흐름, 이의신청서 생성 노드, 통합 QA"],
                ["leejaegang27", "경위서/OCR 결과 처리, 텍스트 ML, 과실비율 판례, 유튜브 자막 사례, 과실비율심의사례 데이터, 판례 Agent"],
                ["ohjuheecode", "차량 사고 이미지·영상 데이터셋, Vision/DL 분석, 영상·이미지 Agent, DL 결과 구조화"],
                ["techshin31", "법률 데이터 수집, 전처리, DB 적재, 법률 원문/조문/근거 metadata"],
                ["workzion2", "고지서 OCR, 과태료·범칙금·벌칙 분석용 룰/매핑 데이터, 과태료·범칙금 분석 흐름"],
            ],
            [45 * mm, 126 * mm],
        )
    )

    doc = SimpleDocTemplate(
        str(PROJECT_PDF),
        pagesize=A4,
        rightMargin=17 * mm,
        leftMargin=17 * mm,
        topMargin=16 * mm,
        bottomMargin=18 * mm,
        title="프로젝트 기획서",
        author="SKN27-FINAL-3Team",
        subject="프로젝트 기획서",
    )
    doc.build(story, onFirstPage=portrait_footer("프로젝트 기획서"), onLaterPages=portrait_footer("프로젝트 기획서"))
    return PROJECT_PDF


def build_data_report_pdf() -> Path:
    story: list = [
        pp("수집 데이터 보고서", "title"),
        pp(f"{TEAM_LABEL}\n프로젝트명: {PROJECT_NAME}\n작성일: {WRITE_DATE}", "subtitle"),
        ptable(
            [
                ["항목", "내용"],
                ["산출물 단계", "데이터 수집 및 저장"],
                ["깃허브 경로", GITHUB_URL],
                ["작성 팀원", AUTHORS],
                ["주의", "저장소 내 raw 데이터 파일이 확인되지 않은 항목은 진행 중 또는 검증 필요로 표기했다."],
            ],
            [35 * mm, 202 * mm],
        ),
        pp("1. 수집 데이터 개요", "h1"),
        ptable(
            [
                ["데이터명", "수집 대상", "수집 목적", "사용 예정 기능", "출처/저작권"],
                ["법률 원문 데이터", "도로교통법, 시행령, 시행규칙, 행정 기준, 고시", "과태료·범칙금 분석과 법률 근거 검색", "법률 근거 검색 노드, RAG evidence", "공개 법령/API, 출처 metadata 필요"],
                ["과태료·범칙금 룰/매핑", "위반 유형, 금액, 벌점, 예외 조건, 처분 단계", "고지서 분석 결과 구조화", "고지서 OCR·과태료/범칙금 분석 노드", "법령 기반 내부 구조화 데이터"],
                ["판례·심의·사고 사례 텍스트", "판례, 과실비율심의사례, 유튜브 자막 사고 사례", "사고 유형, 쟁점, 유사 사례 검색", "텍스트 ML/판례·사례 검색 노드", "공개 원문/자막, 사용권 검토 필요"],
                ["영상·이미지 manifest", "사고 사진, 차량 파손 이미지, 블랙박스/사고 영상, key frame", "Vision/DL 분석과 장면 요약", "영상·이미지 분석 노드", "비식별 샘플만 저장, 원본 민감 파일 커밋 금지"],
                ["사용자 업로드 자료", "고지서, 경위서, 사고 사진, 영상, 사용자 설명", "개별 사건 분석 입력", "챗봇, 리포트, 이의신청서 초안", "사용자 제공 자료, 개인정보 보호 대상"],
            ],
            [36 * mm, 55 * mm, 55 * mm, 55 * mm, 36 * mm],
        ),
        pp("2. 수집 방법 및 자동화 절차", "h1"),
        ptable(
            [
                ["항목", "내용"],
                ["수집 방식", "API 호출, 웹 크롤링 일부, 사용자 입력, 문서/이미지/영상 업로드, 회의/이슈/요구사항 문서 기반 데이터 범위 정리"],
                ["사용 언어/라이브러리", "Python requests/BeautifulSoup, OCR 도구 후보, RAG chunking/embedding 도구, Django API. 실제 라이브러리는 구현 문서 확정 후 고정"],
                ["자동화 여부 및 주기", "현재는 설계 단계. 법률/룰/사례/영상 manifest는 ETL 파이프라인으로 자동화 예정이며 샘플 검증 전까지 수동 검토 병행"],
                ["오류 처리", "OCR 실패, 필수 필드 누락, RAG 검색 실패, Agent 응답 실패, 파일 업로드 실패를 status=partial/failed와 limitations로 반환"],
                ["흐름", "원천 데이터 확인 -> 수집 및 전처리 -> 저장 및 색인 -> Agent evidence/limitations/next_actions 연결"],
            ],
            [42 * mm, 195 * mm],
        ),
        pp("3. 데이터 설명 및 구성", "h1"),
        ptable(
            [
                ["파일명/테이블명", "필드명", "데이터 타입", "설명", "예시"],
                ["legal_documents", "law_name", "string", "법령명", "도로교통법"],
                ["legal_documents", "article", "string", "조문 번호", "제32조"],
                ["legal_documents", "effective_date", "date", "시행일 또는 기준일", "2026-06-22"],
                ["fine_rules", "violation_type", "string", "위반 유형", "어린이보호구역 주정차 위반"],
                ["fine_rules", "fine_amount", "integer", "과태료 또는 범칙금 금액", "120000"],
                ["fault_cases", "issue_tags", "array", "과실 쟁점 태그", "신호, 좌회전, 선진입"],
                ["media_manifest", "file_type", "string", "이미지/영상 유형", "blackbox_video"],
                ["agent_results", "evidence", "array", "근거 목록", "law, precedent, vision_result"],
            ],
            [45 * mm, 40 * mm, 30 * mm, 72 * mm, 50 * mm],
        ),
        ptable(
            [
                ["항목", "내용"],
                ["전체 수집 데이터 건수", "문서 기준 법령/시행령/시행규칙 총 23개 수집 완료로 기록됨. 저장소 내 raw 데이터 파일은 확인되지 않아 원본 적재 건수는 검증 필요."],
                ["추출된 고품질 데이터 건수", "법률 데이터 외 판례, 유튜브 자막, 과실비율심의사례, 영상·이미지 샘플은 진행 중 또는 검증 필요 상태로 기록."],
                ["저장 경로", "설계 기준: etl/legal, etl/fine_rules, etl/fault_cases, etl/vision_manifest, storage/rag, storage/samples. 현재 raw 데이터 저장 경로는 미확정."],
                ["저장 포맷", "CSV, JSON, DB table, RAG document/chunk 모두 후보. 구현 전 schema 확정 필요."],
                ["인코딩", "UTF-8 고정. HTML, Markdown, CSV/JSON, DB 입출력 모두 한글 깨짐 방지 기준 적용."],
            ],
            [45 * mm, 192 * mm],
        ),
        pp("4. 법적·윤리적 검토", "h1"),
        ptable(
            [
                ["항목", "검토 내용"],
                ["개인정보 포함 여부", "법령/공개 사례 데이터는 미포함 가능성이 높으나 사용자 업로드 고지서/영상/사진은 개인정보 또는 민감정보 포함 가능."],
                ["포함 가능 필드", "이름, 차량번호, 주소, 연락처, 위반 일시/장소, 영상/사진 속 개인 식별 정보."],
                ["비식별화", "원본 민감 데이터는 Git 커밋 금지. 테스트 fixture는 비식별화 또는 샘플 대체 후 사용."],
                ["출처 및 사용권", "공개 법령은 출처와 최신성 metadata를 남긴다. 판례/자막/영상은 원문 링크와 사용권 검토 필요."],
                ["라이선스/약관", "진행 중. 유튜브 자막, 보험사 자료, 영상 데이터는 사용 범위 확인 필요."],
            ],
            [45 * mm, 192 * mm],
        ),
        pp("5. 데이터 품질 및 정합성 관리 방안", "h1"),
        ptable(
            [
                ["항목", "관리 방안"],
                ["중복 제거 기준", "source_url, law_name+article, case_id, video_id+timestamp_range, file_hash 기준 중복 제거."],
                ["정합성 검증", "필수 필드 누락, 날짜 형식, 금액/벌점 타입, source_type/domain, evidence metadata를 검증."],
                ["Null 처리", "필수 입력이 없으면 partial 상태와 missing_fields를 반환하고 분석보다 추가 질문을 우선한다."],
                ["표준화", "HTML 제거, 광고/불필요 문구 제거, 문단 분리, chunk 분할, metadata 정규화, UTF-8 저장."],
            ],
            [45 * mm, 192 * mm],
        ),
        pp("6. 변경 이력 및 보완 내역", "h1"),
        ptable(
            [
                ["변경일", "변경자", "변경 내용", "비고"],
                ["2026-06-17", "3팀", "요구사항 정의서 v0.4 업데이트", "125개 요구사항으로 재정리"],
                ["2026-06-18", "3팀", "WBS, 역할, Agent/Supervisor 구조 정리", "중간/최종 일정 기준 반영"],
                ["2026-06-22", "3팀", "수집 데이터 보고서 PDF 작성", "raw 데이터 미확인 항목은 검증 필요로 표기"],
            ],
            [32 * mm, 32 * mm, 115 * mm, 58 * mm],
        ),
    ]

    doc = SimpleDocTemplate(
        str(DATA_PDF),
        pagesize=landscape(A4),
        rightMargin=12 * mm,
        leftMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=16 * mm,
        title="수집 데이터 보고서",
        author="SKN27-FINAL-3Team",
        subject="수집 데이터 보고서",
    )
    doc.build(story, onFirstPage=landscape_footer("수집 데이터 보고서"), onLaterPages=landscape_footer("수집 데이터 보고서"))
    return DATA_PDF


def screen_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont(PDF_FONT, 7)
    canvas.setFillColor(colors.HexColor("#6B7280"))
    canvas.drawString(doc.leftMargin, 8 * mm, "교통분쟁 AI 서비스 화면설계서 v0.3")
    canvas.drawRightString(landscape(A4)[0] - doc.rightMargin, 8 * mm, str(doc.page))
    canvas.restoreState()


def build_screen_pdf() -> Path:
    story: list = [
        pp("교통분쟁 AI 서비스 화면설계서", "title"),
        pp(f"{TEAM_LABEL}\n프로젝트명: {PROJECT_NAME}\n작성일: {WRITE_DATE} / 문서 버전: v0.3\n기준 문서: 요구사항 정의서 v0.4, docs/screen-design-specification.md, screen-design-mvp-flow.html", "subtitle"),
        pp("1. 문서 목적", "h1"),
        pp("본 문서는 교통분쟁 AI 서비스의 MVP 화면 구조, 화면별 표시 데이터, 사용자 액션, 진입 경로, 요구사항 매핑을 정의한다. 화면설계서 양식은 별도 제공되지 않았으므로 기존 Markdown 화면설계서와 화면 이미지 자산을 기준으로 임의 구성했다."),
        pp("2. 전체 사용자 흐름", "h1"),
        ptable(
            [
                ["순서", "화면", "사용자 목적", "주요 액션", "다음 화면"],
                ["1", "UI-ENTRY-001 로그인/서비스 설명 진입", "서비스 목적을 이해하고 상담 시작", "Google 로그인 또는 챗봇 바로 시작", "UI-AUTH-001 또는 UI-Ai-01"],
                ["2", "UI-AUTH-001 Google 로그인", "이력과 리포트 저장 가능 상태 진입", "Google 로그인, 약관 동의", "UI-Ai-01"],
                ["3", "UI-Ai-01 AI 교통 상담 챗봇", "고지서, 사고 설명, 사진·영상, 법령 질문 입력", "텍스트 입력, 파일 업로드, 추천 질문 선택", "Supervisor 분기"],
                ["4", "Supervisor 분기", "입력 유형에 맞는 내부 노드 호출", "사용자는 직접 노드를 선택하지 않음", "결과 카드"],
                ["5", "결과/리포트", "분석 요약, 근거, 한계, 후속 행동 확인", "상세보기, 근거 보기, 리포트 생성", "마이페이지, 리포트 상세"],
            ],
            [12 * mm, 50 * mm, 70 * mm, 62 * mm, 50 * mm],
        ),
        pp("3. 화면 목록", "h1"),
        ptable(
            [
                ["Screen ID", "화면명", "설명", "MVP 포함"],
                ["UI-ENTRY-001", "로그인/서비스 설명 진입", "서비스 목적, Google 로그인, 챗봇 바로 시작 CTA 제공", "포함"],
                ["UI-AUTH-001", "Google 로그인", "Google 로그인 및 신규 사용자 약관 동의 흐름", "포함"],
                ["UI-MY-001", "마이페이지", "분석 현황, 기한 임박 사건, 최근 분석 이력 요약", "포함"],
                ["UI-Ai-01", "AI 교통 상담 챗봇", "교통사고, 과실비율, 범칙금/과태료 상담 대화와 분석 카드", "포함"],
                ["UI-HIS-001", "과거 이력", "분석 이력 목록, 유형/기간/검색 필터, 상세 진입", "포함"],
                ["UI-REPORT-FAULT-001", "사고 과실비율 분석 리포트", "사고 개요, AI 분석 결과, 판단 근거, 유사 사례, 후속 조치", "포함"],
                ["UI-REPORT-FINE-001", "과태료·범칙금 대응 리포트", "OCR, 처분 결과, 이의제기 가능성, 필요 증거, 법령/판례, 초안", "포함"],
            ],
            [35 * mm, 48 * mm, 130 * mm, 24 * mm],
        ),
        pp("4. 공통 UI 기준", "h1"),
        ptable(
            [
                ["구분", "기준"],
                ["인코딩", "모든 문서와 프론트엔드 소스는 UTF-8로 저장하고 HTML에는 meta charset=UTF-8을 명시한다."],
                ["언어/폰트", "HTML 루트 lang=ko, Pretendard/Noto Sans KR/Apple SD Gothic Neo/Malgun Gothic/sans-serif 권장."],
                ["상태 처리", "로딩, 빈 데이터, 오류, 권한 없음, 데이터 수집 진행 중 상태를 화면별로 제공한다."],
                ["주의 문구", "AI 분석은 참고용이며 실제 판단과 다를 수 있음을 리포트와 결과 화면에 표시한다."],
                ["다운로드", "리포트와 이의신청서 초안은 PDF/DOCX 저장 또는 다운로드 액션을 제공한다."],
            ],
            [34 * mm, 203 * mm],
        ),
        pp("5. 주요 화면 상세", "h1"),
        ptable(
            [
                ["화면", "핵심 구성", "주요 액션", "예외/빈 상태"],
                ["마이페이지", "요약 카드, 기한 임박 사건, 최근 분석 이력, 과거이력 이동", "상세보기, 과거이력 이동", "등록 사건 없음, 기한 임박 없음, 조회 실패"],
                ["챗봇", "대화 목록, 메시지, 분석 카드, 근거 버튼, 파일 첨부, 입력창", "새 대화, 상세 보기, 근거 보기, 파일 첨부", "RAG 근거 부족, 업로드 실패, 응답 실패"],
                ["과실비율 리포트", "사고 개요, 제출 자료, AI 분석, 판단 근거, 핵심 쟁점, 유사 판례", "유사 판례 보기, 리포트 저장", "자료 부족, 과실비율 단정 금지"],
                ["과태료·범칙금 리포트", "OCR 결과, 처분 단계, 이의제기 가능성, 필요 증거, 법령/판례, 초안", "이미지 확대, 법령 더보기, 초안 복사, PDF 저장", "OCR 실패, 신뢰도 낮음, 초안 생성 실패"],
            ],
            [40 * mm, 83 * mm, 58 * mm, 56 * mm],
        ),
        pp("6. API 연결 초안", "h1"),
        ptable(
            [
                ["화면", "필요 데이터", "API 정의서에서 확정할 항목"],
                ["UI-MY-001", "마이페이지 요약, 기한 임박 사건, 최근 분석 이력", "GET /api/mypage/summary/, GET /api/history/"],
                ["UI-Ai-01", "세션 목록, 메시지, 챗봇 응답, 분석 카드, 첨부 파일", "GET /api/chat/sessions/, POST /api/chat/*"],
                ["UI-REPORT-FAULT-001", "사고 개요, 제출 자료, 분석 결과, 판단 근거, 유사 판례", "GET /api/reports/{id}/ 또는 과실비율 전용 endpoint"],
                ["UI-REPORT-FINE-001", "OCR 결과, 처분 결과, 이의제기 가능성, 법령/판례, 초안", "GET /api/reports/{id}/ 또는 과태료·범칙금 전용 endpoint"],
            ],
            [48 * mm, 95 * mm, 94 * mm],
        ),
        pp("7. 남은 결정 작업", "h1"),
        ptable(
            [
                ["우선순위", "항목", "남은 작업"],
                ["높음", "Screen ID 체계", "화면정의서 이미지의 ID와 요구사항 정의서 ID 체계 매핑 확정"],
                ["높음", "진단하기", "과실비율 진단과 과태료·범칙금 진단 입력 폼 분리 여부 확정"],
                ["높음", "API 정의서", "과실비율 상세, 리포트, 과태료·범칙금 리포트 endpoint 분리 기준 확정"],
                ["보통", "챗봇", "범칙금 챗봇과 보험 챗봇을 단일 화면에서 구분하는 방식 확정"],
                ["보통", "판례/보험사 사례 목록", "수집 진행 중, 원문 링크 없음, 결과 없음 상태 표시 기준 확정"],
            ],
            [22 * mm, 55 * mm, 160 * mm],
        ),
        PageBreak(),
        pp("8. 화면 자산", "h1"),
    ]

    asset_dir = ROOT / "docs" / "assets" / "screen-design"
    add_pdf_image(story, asset_dir / "final-screen-plan-complete-updated.png", "통합 화면정의서 이미지", 245 * mm, 135 * mm)
    add_pdf_image(story, asset_dir / "mypage-updated.png", "마이페이지 업데이트 화면", 245 * mm, 130 * mm)
    add_pdf_image(story, asset_dir / "report-pages-updated.png", "과실비율 리포트 및 과태료·범칙금 대응 리포트 화면", 245 * mm, 130 * mm)

    doc = SimpleDocTemplate(
        str(SCREEN_PDF),
        pagesize=landscape(A4),
        rightMargin=12 * mm,
        leftMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=16 * mm,
        title="교통분쟁 AI 서비스 화면설계서",
        author="SKN27-FINAL-3Team",
        subject="화면설계서",
    )
    doc.build(story, onFirstPage=screen_footer, onLaterPages=screen_footer)
    return SCREEN_PDF


def main() -> None:
    ensure_dirs()
    project_docx = PROJECT_DOCX
    data_docx = DATA_DOCX
    if PROJECT_TEMPLATE.exists():
        project_docx = build_project_plan()
    elif not PROJECT_DOCX.exists():
        project_docx = None

    if DATA_TEMPLATE.exists():
        data_docx = build_data_report()
    elif not DATA_DOCX.exists():
        data_docx = None

    project_pdf = build_project_plan_pdf()
    data_pdf = build_data_report_pdf()
    screen_pdf = build_screen_pdf()

    if project_docx:
        print(project_docx)
    if data_docx:
        print(data_docx)
    print(project_pdf)
    print(data_pdf)
    print(screen_pdf)


if __name__ == "__main__":
    main()
