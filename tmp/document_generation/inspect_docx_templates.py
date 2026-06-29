from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def compact(text: str, limit: int = 160) -> str:
    normalized = " ".join(text.split())
    if len(normalized) <= limit:
        return normalized
    return normalized[: limit - 3] + "..."


def inspect(path: Path) -> None:
    doc = Document(path)
    print(f"FILE\t{path}")
    print(f"PARAGRAPHS\t{len(doc.paragraphs)}")
    for index, paragraph in enumerate(doc.paragraphs[:80], start=1):
        text = compact(paragraph.text)
        if text:
            style = paragraph.style.name if paragraph.style else ""
            print(f"P{index:03d}\t{style}\t{text}")

    print(f"TABLES\t{len(doc.tables)}")
    for table_index, table in enumerate(doc.tables, start=1):
        row_count = len(table.rows)
        col_count = len(table.columns) if table.rows else 0
        print(f"T{table_index:03d}\t{row_count}x{col_count}")
        for row_index, row in enumerate(table.rows[:8], start=1):
            cells = [compact(cell.text, 80) for cell in row.cells]
            print(f"T{table_index:03d}R{row_index:02d}\t" + " | ".join(cells))


def main() -> int:
    for arg in sys.argv[1:]:
        inspect(Path(arg))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
